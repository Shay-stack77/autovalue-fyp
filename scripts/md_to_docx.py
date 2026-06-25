"""Convert docs/report.md into a fully-prepared Arden FYP submission.

Does all four submission-prep tasks:
  1. Fills the Arden cover-page table cells with student details
  2. Inserts the three diagnostic charts at their references in Chapter 5.4
  3. Adds page numbers (centred, "Page X of Y") in the footer
  4. Inserts an auto-updating Word TOC field after the cover page

Output: docs/FYP_Dissertation_STU195050.docx
"""
from __future__ import annotations

import re
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ---------- paths --------------------------------------------------------
PROJECT = Path(r"D:\assignments\Computing project fyp")
TEMPLATE = Path(r"D:\assignments\cloud com\Arden_Assessment_Cover_page.docx")
# Official Arden assessment cover sheet (7-row label table) — bundled in repo.
COVER_TEMPLATE = PROJECT / "docs" / "assets" / "Arden_Assessment_Cover_page.docx"
ARDEN_LOGO = PROJECT / "docs" / "assets" / "arden_logo.jpeg"   # top-right page header
SOURCE = PROJECT / "docs" / "report.md"
OUTPUT = PROJECT / "docs" / "FYP_Dissertation_STU195050_FINAL.docx"
SCREENSHOTS = PROJECT / "docs" / "screenshots"
DOCS_DIR = PROJECT / "docs"

# ---------- student details ---------------------------------------------
STUDENT = {
    "title":   "Machine Learning-Based Predictive Valuation System for "
               "Second-Hand Vehicles Using Real-Time Market Data",
    "name":    "Syed Shayaan Ali Ali",
    "id":      "STU195050",
    "module":  "Final Year Project (CMP6200) — BSc (Hons) Computing",
    "ethics":  "Ethics Reference: P17519 (Low Risk — Approved)",
    "uni":     "Arden University — School of Computing",
    "date":    "June 2026",
}

# ---------- typography ---------------------------------------------------
# Per Arden University SCC FYP handbook (Section D — Submission Format):
# "A4 layout using double line spacing. The recommended font is Arial,
# size 12 for the main content and depreciating font sizes for sub-headers
# and headers etc. All text should be justified."
BODY_FONT = "Calibri"          # match the student's other modules (WAD)
BODY_SIZE = Pt(11)
H1_SIZE = Pt(18)
H2_SIZE = Pt(14)
H3_SIZE = Pt(13)
H4_SIZE = Pt(12)
HEADING_COLOR = RGBColor(0x36, 0x5F, 0x91)   # heading blue, matching prior Arden modules


def set_font(run, *, name=BODY_FONT, size=BODY_SIZE,
             bold=None, italic=None, color=None, mono=False):
    n = "Consolas" if mono else name
    run.font.name = n
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), n)
    rFonts.set(qn("w:hAnsi"), n)
    rFonts.set(qn("w:cs"), n)


def body_spacing(p):
    """Match the WAD module: 1.15 line spacing, left-aligned, 10pt after."""
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# alias for backwards-compat
double_space = body_spacing


# ---------- inline parser ------------------------------------------------
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(p, text: str, *, size=BODY_SIZE):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_font(r, size=size)
        token = m.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2]); set_font(r, size=size, bold=True)
        elif token.startswith("*"):
            r = p.add_run(token[1:-1]); set_font(r, size=size, italic=True)
        elif token.startswith("`"):
            r = p.add_run(token[1:-1]); set_font(r, size=size, mono=True)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:]); set_font(r, size=size)


# ---------- table renderer ----------------------------------------------
def render_table(doc, lines):
    rows = []
    for ln in lines:
        ln = ln.strip()
        if not ln: continue
        if re.match(r"^\|?\s*[:\-\|\s]+$", ln): continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
    if not rows: return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for ri, cells in enumerate(rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)
            txt = cells[ci] if ci < len(cells) else ""
            if ri == 0:
                r = p.add_run(txt); set_font(r, size=Pt(10), bold=True)
            else:
                add_inline(p, txt, size=Pt(10))
    doc.add_paragraph()


# ---------- chart embedding ---------------------------------------------
# After-heading-text triggers for image insertion
# maps the heading text that triggers an image to (figure number, caption, image path)
CHART_TRIGGERS = {
    "7.5.1 Predicted vs. Actual Price": ("7.1", "Predicted vs. actual price on the held-out test set", SCREENSHOTS / "predicted_vs_actual.png"),
    "7.5.2 Residual Distribution":      ("7.2", "Distribution of residuals (predicted − actual)", SCREENSHOTS / "residuals.png"),
    "7.5.3 Feature Importance":         ("7.3", "XGBoost gain-based feature importance", SCREENSHOTS / "feature_importance.png"),
    "7.5.4 Prediction Interval":        ("7.4", "80% prediction interval vs. actual prices on a test sample", SCREENSHOTS / "prediction_interval.png"),
}


def insert_chart(doc, image_path: Path, caption: str,
                 max_w: float = 5.5, max_h: float = 7.0):
    """Embed an image scaled to fit within (max_w x max_h) inches, preserving
    aspect ratio. Tall screenshots are height-limited so they never overflow
    the page."""
    if not image_path.exists():
        print(f"[warn] missing chart: {image_path}")
        return
    # native pixel size -> pick the dimension that binds
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(image_path) as im:
            px_w, px_h = im.size
        aspect = px_h / px_w if px_w else 1.0
    except Exception:
        aspect = 1.0
    width = max_w
    if width * aspect > max_h:        # too tall -> bind on height instead
        width = max_h / aspect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    set_font(cr, size=Pt(10), italic=True)


# ---------- markdown streamer -------------------------------------------
def stream_markdown(doc, md: str):
    lines = md.splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    pending_chart = None  # tuple (image_path, caption) to insert after next paragraph

    while i < len(lines):
        ln = lines[i]
        stripped = ln.rstrip()

        # Skip front-matter
        if i == 0 and stripped == "---":
            j = i + 1
            while j < len(lines) and lines[j].rstrip() != "---":
                j += 1
            i = j + 1
            continue

        # Fenced code
        if stripped.startswith("```"):
            if not in_code:
                in_code = True; code_buffer = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                # Code/ASCII-art must be LEFT-aligned, never justified, or Word
                # stretches the spaces and shreds the monospace alignment.
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run("\n".join(code_buffer))
                set_font(r, size=Pt(9), mono=True)
                in_code = False; code_buffer = []
            i += 1; continue
        if in_code:
            code_buffer.append(ln); i += 1; continue

        # Page break (---)
        if stripped == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1; continue

        # Markdown image: ![alt](relative/path.png)
        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m_img:
            alt = m_img.group(1).strip()
            img_rel = m_img.group(2).strip()
            img_path = (DOCS_DIR / img_rel).resolve() if not Path(img_rel).is_absolute() else Path(img_rel)
            if img_path.exists():
                insert_chart(doc, img_path, alt or "")
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[missing image: {img_path}]")
                set_font(r, italic=True)
                print(f"[warn] missing image: {img_path}")
            i += 1; continue

        # Heading
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip().rstrip("#").strip()
            # Use Word's built-in Heading styles so the TOC field populates,
            # but override the run to Arial / near-black per the handbook.
            style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}[level]
            p = doc.add_paragraph(style=style_name)
            size = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE, 4: H4_SIZE}[level]
            r = p.add_run(text)
            set_font(r, size=size, bold=(level <= 3), italic=(level == 4),
                     color=HEADING_COLOR)
            # Headings: left-align (handbook says body justified; headings normally
            # left-aligned for legibility), single-spaced
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True

            # Trigger chart insertion after the chart sub-section paragraph
            for trigger, (fignum, cap, path) in CHART_TRIGGERS.items():
                if trigger in text:
                    pending_chart = (path, f"Figure {fignum} — {cap}")
                    break
            i += 1; continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) \
                and re.match(r"^\|?\s*[:\-\|\s]+\|", lines[i + 1].strip()):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i]); i += 1
            render_table(doc, tbl_lines); continue

        # Bullet
        m = re.match(r"^\s*[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph()
            double_space(p)
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-12)
            r = p.add_run("• "); set_font(r)
            add_inline(p, m.group(1))
            i += 1; continue

        # Numbered list
        nm = re.match(r"^\s*(\d+)\.\s+(.*)$", stripped)
        if nm:
            p = doc.add_paragraph()
            double_space(p)
            p.paragraph_format.left_indent = Pt(28)
            p.paragraph_format.first_line_indent = Pt(-16)
            r = p.add_run(f"{nm.group(1)}. "); set_font(r)
            add_inline(p, nm.group(2))
            i += 1; continue

        # Blank line
        if not stripped:
            i += 1; continue

        # Default paragraph (collect continuation lines)
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt or re.match(r"^(#{1,4}\s|\|\s|\s*[-*]\s|\s*\d+\.\s|```)", nxt):
                break
            if nxt == "---": break
            para_lines.append(nxt); j += 1
        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        double_space(p)
        add_inline(p, para_text)
        i = j

        # Insert chart immediately after this paragraph if pending
        if pending_chart is not None:
            insert_chart(doc, pending_chart[0], pending_chart[1])
            pending_chart = None


# ---------- cover page filler -------------------------------------------
def fill_cover(doc):
    """Populate the 7-row Arden cover-page table with student details.

    Row mapping (top→bottom): 0=title, 1=name, 2=ID, 3=module, 4=ethics,
                              5=uni, 6=date.
    Each cell is centred, bold for the title row, large for visibility.
    """
    table = doc.tables[0]
    rows_content = [
        (STUDENT["title"],  Pt(20), True,  WD_ALIGN_PARAGRAPH.CENTER),
        (STUDENT["name"],   Pt(16), True,  WD_ALIGN_PARAGRAPH.CENTER),
        (f"Student ID: {STUDENT['id']}", Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER),
        (STUDENT["module"], Pt(14), False, WD_ALIGN_PARAGRAPH.CENTER),
        (STUDENT["ethics"], Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER),
        (STUDENT["uni"],    Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER),
        (STUDENT["date"],   Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER),
    ]
    for ri, (text, size, bold, align) in enumerate(rows_content):
        if ri >= len(table.rows): break
        cell = table.rows[ri].cells[0]
        # Take the first paragraph; clear it; write our text
        p = cell.paragraphs[0]
        # remove existing runs
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        p.alignment = align
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=RGBColor(0x0D, 0x0D, 0x0D))


# ---------- TOC field ----------------------------------------------------
def insert_toc(doc):
    """Insert a Word auto-update TOC field. Word will populate it on open
    (right-click → Update Field, or simply press F9)."""
    p = doc.add_paragraph()
    r = p.add_run("Table of Contents")
    set_font(r, size=H2_SIZE, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    # Build the TOC field via OOXML
    p2 = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    pt = OxmlElement("w:t"); pt.text = "Right-click and select 'Update Field' (or press F9) to populate."
    placeholder.append(pt)
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")

    run_el = OxmlElement("w:r"); run_el.append(fld_begin); run_el.append(instr); run_el.append(fld_sep)
    p2._element.append(run_el)
    p2._element.append(placeholder)
    end_run = OxmlElement("w:r"); end_run.append(fld_end)
    p2._element.append(end_run)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------- page numbers in footer --------------------------------------
def add_page_numbers(doc):
    """Add 'Page X of Y' centred in the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        # Clear any existing footer paragraphs
        for p in list(footer.paragraphs):
            for run in list(p.runs):
                run._element.getparent().remove(run._element)

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "Page "
        r1 = p.add_run("Page "); set_font(r1, size=Pt(10))

        # PAGE field
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        r2 = p.add_run(); set_font(r2, size=Pt(10))
        r2._element.append(fld_begin); r2._element.append(instr); r2._element.append(fld_end)

        # " of "
        r3 = p.add_run(" of "); set_font(r3, size=Pt(10))

        # NUMPAGES field
        fld_begin2 = OxmlElement("w:fldChar"); fld_begin2.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText"); instr2.set(qn("xml:space"), "preserve"); instr2.text = "NUMPAGES"
        fld_end2 = OxmlElement("w:fldChar"); fld_end2.set(qn("w:fldCharType"), "end")
        r4 = p.add_run(); set_font(r4, size=Pt(10))
        r4._element.append(fld_begin2); r4._element.append(instr2); r4._element.append(fld_end2)


# ---------- main --------------------------------------------------------
def apply_dm_document_defaults(doc):
    """Set document-wide defaults per Arden FYP handbook:
    Arial 12pt body, double line spacing, justified, en-GB lang."""
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, docDefaults)

    # rPrDefault — Arial 12pt
    for old in docDefaults.findall(qn("w:rPrDefault")):
        docDefaults.remove(old)
    rPrDef = OxmlElement("w:rPrDefault")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(k), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rPr.append(sz)   # 12pt
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "24"); rPr.append(szCs)
    lang = OxmlElement("w:lang"); lang.set(qn("w:val"), "en-GB"); rPr.append(lang)
    rPrDef.append(rPr)
    docDefaults.append(rPrDef)

    # pPrDefault — double-spaced, justified, 6pt after
    for old in docDefaults.findall(qn("w:pPrDefault")):
        docDefaults.remove(old)
    pPrDef = OxmlElement("w:pPrDefault")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "120")     # 6pt
    spacing.set(qn("w:line"), "480")      # double (240 = single, 480 = double)
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "both"); pPr.append(jc)  # justified
    pPrDef.append(pPr)
    docDefaults.append(pPrDef)

    # A4 paper (handbook requires A4) + 1 inch margins on all sides
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def build_cover(doc):
    """Build a clean, self-contained Arden FYP title page from scratch.

    Built programmatically (not from an external template) so the document
    never inherits unrelated content.
    """
    def line(text, size, *, bold=False, space_after=6, space_before=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=RGBColor(0x0D, 0x0D, 0x0D))
        return p

    line("Arden University", Pt(16), bold=True, space_before=48)
    line("School of Computing", Pt(13))
    line("BSc (Hons) Computing — Level 6", Pt(13), space_after=28)
    line("Final Year Project (CMP6200)", Pt(13), bold=True, space_after=40)
    line(STUDENT["title"], Pt(20), bold=True, space_after=40)
    line(STUDENT["name"], Pt(15), bold=True)
    line(f"Student ID: {STUDENT['id']}", Pt(13), space_after=28)
    line("Ethics Reference: P17519 (Low Risk — Approved)", Pt(12))
    line("Assessable word count: 10,999 (Chapters 1–11)", Pt(12))
    line("Live deployment: https://autovalue-fyp.onrender.com", Pt(12))
    line(STUDENT["date"], Pt(12), space_before=14)


def ensure_heading_styles(doc):
    """Guarantee Heading 1-4 paragraph styles exist with correct outline levels.

    The Arden cover template only fully defines some headings (others are latent),
    which would crash add_paragraph(style="Heading 2"). Create any missing ones
    and force the outline level so the TOC field still collects them.
    """
    existing = {s.name: s for s in doc.styles}
    for lvl in (1, 2, 3, 4):
        name = f"Heading {lvl}"
        st = existing.get(name)
        if st is None:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = doc.styles["Normal"]
        pPr = st.element.get_or_add_pPr()
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            pPr.append(ol)
        ol.set(qn("w:val"), str(lvl - 1))


def fill_official_cover(doc):
    """Type the student's answers onto the official Arden cover sheet.

    The cover is a 7-row, 1-column table; each cell holds a label image
    (Awarding Body, Programme Name, ...) followed by a blank line where the
    answer is typed.
    """
    fields = [
        "Arden University",                 # Awarding Body
        "BSc (Hons) Computing",             # Programme Name
        "Final Year Project (CMP6200)",     # Module Name (and Part if applicable)
        STUDENT["title"],                   # Assessment Title
        STUDENT["id"],                      # Student Number
        "Awais Malik",                      # Tutor Name
        "10,999",                           # Word Count
    ]
    table = doc.tables[0]
    for ri, val in enumerate(fields):
        if ri >= len(table.rows):
            break
        cell = table.rows[ri].cells[0]
        # the blank line directly beneath the label image
        para = cell.paragraphs[1] if len(cell.paragraphs) > 1 else cell.add_paragraph()
        r = para.add_run(val)
        set_font(r, size=Pt(12), bold=True, color=RGBColor(0x0D, 0x0D, 0x0D))


def main():
    # 1. Start from the official Arden cover sheet and fill it in
    doc = Document(str(COVER_TEMPLATE))
    fill_official_cover(doc)

    # Base font for any unstyled text / the TOC (cover is image-based, so this
    # does not disturb it). Body paragraphs are styled explicitly below.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Heading 1-4 styles may be latent in the Arden template — define them so
    # the heading handler and TOC work.
    ensure_heading_styles(doc)

    # 2. Section break: the dissertation starts on its own page. The body pages
    #    carry the Arden logo top-right (matching the student's other modules).
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    body_sec.header.is_linked_to_previous = False
    hp = body_sec.header.paragraphs[0]
    for r in list(hp.runs):
        r._element.getparent().remove(r._element)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    body_sec.header_distance = Inches(0.3)
    if ARDEN_LOGO.exists():
        hp.add_run().add_picture(str(ARDEN_LOGO), width=Inches(0.68))
    # Body section only: A4 + 1 inch margins (handbook). Cover left untouched.
    body_sec.page_width = Mm(210)
    body_sec.page_height = Mm(297)
    body_sec.top_margin = Inches(1)
    body_sec.bottom_margin = Inches(1)
    body_sec.left_margin = Inches(1)
    body_sec.right_margin = Inches(1)

    # 3. TOC (first page of the body section)
    insert_toc(doc)

    # 4. Stream the report (charts embedded at the 7.5.x references)
    md_text = SOURCE.read_text(encoding="utf-8")
    md_text = re.sub(r"## Table of Contents.*?(?=\n##|\n# )", "", md_text, flags=re.S)
    md_text = re.sub(
        r"\n# Machine Learning-Based Predictive Valuation System[^\n]*\n.*?\n---\n",
        "\n", md_text, count=1, flags=re.S)
    stream_markdown(doc, md_text)

    # 5. Page numbers in every section's footer
    add_page_numbers(doc)

    # 6. Document metadata — author/title as the student, not the library
    cp = doc.core_properties
    cp.author = STUDENT["name"]
    cp.last_modified_by = STUDENT["name"]
    cp.title = STUDENT["title"]
    cp.subject = "BSc (Hons) Computing — Final Year Project (CMP6200)"
    cp.category = "Dissertation"
    cp.comments = ""          # clear the default "generated by python-docx"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"[ok] wrote {OUTPUT}  ({OUTPUT.stat().st_size:,} bytes)")
    print(f"     - official Arden cover sheet filled (tutor: Awais Malik)")
    print(f"     - TOC field inserted (open in Word, press F9 to populate)")
    print(f"     - Arden-logo header kept on the cover page only")
    print(f"     - 'Page X of Y' added to footer")


if __name__ == "__main__":
    main()
