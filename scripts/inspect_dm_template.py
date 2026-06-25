"""Inspect the user's prior DM assignment as a style template."""
from docx import Document
from docx.shared import Pt

path = r"D:\assignments\data mining\STU195050-DM.docx"
doc = Document(path)

print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Sections: {len(doc.sections)}")

# Page setup
s = doc.sections[0]
print(f"\nPage size: {s.page_width.inches:.2f} x {s.page_height.inches:.2f} inches")
print(f"Margins: top={s.top_margin.inches:.2f} bottom={s.bottom_margin.inches:.2f} "
      f"left={s.left_margin.inches:.2f} right={s.right_margin.inches:.2f}")

# Default style
normal = doc.styles["Normal"]
print(f"\nNormal style font: name={normal.font.name!r}, size={normal.font.size}")

# Sample paragraphs — show first 40 with style and run formatting
print("\n--- First 50 non-empty paragraphs ---")
shown = 0
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip():
        continue
    style_name = p.style.name if p.style else "?"
    align = p.alignment
    spacing = p.paragraph_format.line_spacing
    space_after = p.paragraph_format.space_after
    first_run = p.runs[0] if p.runs else None
    if first_run:
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        italic = first_run.italic
        color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
    else:
        font_name = font_size = bold = italic = color = None
    text = (p.text[:80] + "…") if len(p.text) > 80 else p.text
    print(f"[{i:3d}] style={style_name!r:25} align={align} "
          f"font={font_name!r} size={font_size} bold={bold} italic={italic} "
          f"color={color} space_after={space_after} line={spacing}")
    print(f"      text={text!r}")
    shown += 1
    if shown >= 50:
        break

# Header / footer
print("\n--- Header ---")
for p in s.header.paragraphs:
    print(f"  {p.text!r}")
print("--- Footer ---")
for p in s.footer.paragraphs:
    print(f"  {p.text!r}")

# Heading styles in use
print("\n--- Distinct paragraph styles in use ---")
seen = {}
for p in doc.paragraphs:
    sn = p.style.name if p.style else "?"
    seen[sn] = seen.get(sn, 0) + 1
for k, v in sorted(seen.items(), key=lambda x: -x[1]):
    print(f"  {k}: {v}")
