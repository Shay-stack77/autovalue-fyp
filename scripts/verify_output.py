"""Verify the v2 output has all 4 enhancements."""
from docx import Document
from docx.oxml.ns import qn

path = r"D:\assignments\Computing project fyp\docs\FYP_Dissertation_STU195050_v2.docx"
doc = Document(path)

# 1. Cover page filled?
print("=== 1. Cover page ===")
for ri, row in enumerate(doc.tables[0].rows):
    txt = row.cells[0].text.strip()
    if txt:
        print(f"  Row {ri}: {txt[:80]!r}")

# 2. Charts embedded?
print("\n=== 2. Embedded images ===")
imgs = doc.part.part_related_by_reltype if hasattr(doc.part, "part_related_by_reltype") else None
img_parts = [p for p in doc.part.related_parts.values()
             if "image" in p.content_type]
print(f"  Total image parts: {len(img_parts)}")
for p in img_parts:
    print(f"    - {p.partname} ({p.content_type})")

# 3. Page numbers?
print("\n=== 3. Footer ===")
for s in doc.sections:
    print(f"  Footer text: {s.footer.paragraphs[0].text!r}")
    has_field = bool(s.footer.element.xpath(".//w:fldChar"))
    print(f"  Footer has field codes: {has_field}")

# 4. TOC field?
print("\n=== 4. TOC field ===")
toc_fields = doc.element.body.xpath(".//w:instrText[contains(text(), 'TOC')]")
print(f"  TOC field instructions found: {len(toc_fields)}")
for f in toc_fields:
    print(f"    - {f.text!r}")

# Word count
words = sum(len(p.text.split()) for p in doc.paragraphs)
words += sum(len(c.text.split())
             for tbl in doc.tables for r in tbl.rows for c in r.cells)
print(f"\nApproximate body word count: {words}")
print(f"\nFile size: {__import__('os').path.getsize(path):,} bytes")
