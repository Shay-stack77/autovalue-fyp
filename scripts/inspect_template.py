"""Deeper inspection of the Arden cover-page template."""
from docx import Document

doc = Document(r"C:\Users\HP\Downloads\Arden_Assessment_Cover_page.docx")

print("=== TABLE 0 detailed ===")
table = doc.tables[0]
for ri, row in enumerate(table.rows):
    for ci, cell in enumerate(row.cells):
        print(f"\n[Row {ri}, Col {ci}]")
        for pi, p in enumerate(cell.paragraphs):
            runs_text = [(r.text, r.bold, r.font.size) for r in p.runs]
            print(f"  Para {pi}: style={p.style.name!r}, text={p.text!r}")
            for ri2, r in enumerate(p.runs):
                print(f"    Run {ri2}: text={r.text!r}, bold={r.bold}, size={r.font.size}")

# Also look for textboxes/shapes in the XML body
from lxml import etree
print("\n=== Searching for textboxes / drawings ===")
body_xml = doc.element.body
ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
      "v": "urn:schemas-microsoft-com:vml",
      "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"}
texts = body_xml.xpath("//w:t", namespaces=ns)
print(f"Found {len(texts)} <w:t> nodes")
for t in texts[:30]:
    if t.text and t.text.strip():
        print(f"  TEXT: {t.text!r}")
